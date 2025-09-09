from docx import Document
from docx.shared import Inches
import os
import uuid
import re

class TextEditor:
    def format_text(self, text):
        try:
            # Basic text formatting
            formatted_text = text.strip()
            
            # Remove extra spaces
            formatted_text = re.sub(r'\s+', ' ', formatted_text)
            
            # Fix line breaks
            formatted_text = re.sub(r'\n\s*\n', '\n\n', formatted_text)
            
            # Capitalize sentences
            sentences = formatted_text.split('. ')
            formatted_sentences = []
            for sentence in sentences:
                if sentence.strip():
                    sentence = sentence.strip().capitalize()
                    formatted_sentences.append(sentence)
            
            formatted_text = '. '.join(formatted_sentences)
            
            return formatted_text
        except Exception as e:
            print(f"Text formatting error: {e}")
            return text
    
    def text_to_docx(self, text, output_folder):
        try:
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Extracted Text Document', 0)
            
            # Add text content
            paragraphs = text.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Generate output filename
            output_filename = f"extracted_text_{uuid.uuid4()}.docx"
            output_path = os.path.join(output_folder, output_filename)
            
            # Save document
            doc.save(output_path)
            
            return output_path
        except Exception as e:
            print(f"Text to DOCX error: {e}")
            return None